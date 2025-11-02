"""DOCX resume parser using python-docx.

Extracts text and structure from Word documents (.docx format).
"""

import re
from typing import Dict, List, Optional
from pathlib import Path
from docx import Document
from autoapply.util.logger import get_logger

logger = get_logger(__name__)


class DOCXParseResult:
    """Structured result from DOCX parsing."""

    def __init__(self) -> None:
        self.raw_text: str = ""
        self.paragraphs: List[str] = []
        self.sections: Dict[str, str] = {}
        self.contact_info: Dict[str, Optional[str]] = {
            "name": None,
            "email": None,
            "phone": None,
            "location": None,
            "linkedin": None,
        }
        self.experiences: List[Dict[str, str]] = []
        self.education: List[Dict[str, str]] = []
        self.skills: List[str] = []
        self.tables: List[List[List[str]]] = []  # Nested lists for table data
        self.confidence: float = 0.0


async def parse_docx_resume(file_path: str | Path) -> DOCXParseResult:
    """Parse a DOCX resume and extract structured information.

    python-docx provides better structure than PDF parsing since
    we have access to actual document structure (paragraphs, tables, runs).

    :param file_path: Path to the DOCX file
    :returns: Structured parsing result
    """
    file_path = Path(file_path)
    result = DOCXParseResult()

    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    logger.info(f"Parsing DOCX resume: {file_path}")

    try:
        doc = Document(file_path)

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                result.paragraphs.append(text)

        result.raw_text = "\n".join(result.paragraphs)

        # Extract tables (some resumes use tables for layout)
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result.tables.append(table_data)

    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"Unable to parse DOCX: {e}")

    if not result.raw_text.strip() and not result.tables:
        raise ValueError("DOCX appears to be empty")

    # Extract structured information (similar to PDF parser)
    _extract_contact_info(result)
    _extract_sections(result)
    _extract_experiences(result)
    _extract_education(result)
    _extract_skills(result)

    # Calculate confidence
    result.confidence = _calculate_confidence(result)

    logger.info(
        f"DOCX parsed: {len(result.experiences)} experiences, "
        f"{len(result.education)} education, confidence={result.confidence:.2f}"
    )

    return result


def _extract_contact_info(result: DOCXParseResult) -> None:
    """Extract contact information from the document."""
    text = result.raw_text

    # Email
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text, re.IGNORECASE)
    if email_match:
        result.contact_info["email"] = email_match.group(0)

    # Phone
    phone_patterns = [
        r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
        r"\+\d{1,3}[-.\s]?\(?\d{2,3}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}",
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            result.contact_info["phone"] = phone_match.group(0)
            break

    # LinkedIn
    linkedin_match = re.search(
        r"linkedin\.com/in/[\w-]+|linkedin\.com/[\w-]+", text, re.IGNORECASE
    )
    if linkedin_match:
        result.contact_info["linkedin"] = linkedin_match.group(0)

    # Name (first non-empty paragraph, heuristic)
    if result.paragraphs:
        first_para = result.paragraphs[0]
        if (
            2 <= len(first_para.split()) <= 4
            and first_para[0].isupper()
            and "@" not in first_para
            and not any(char.isdigit() for char in first_para)
        ):
            result.contact_info["name"] = first_para


def _extract_sections(result: DOCXParseResult) -> None:
    """Split document into major sections."""
    text = result.raw_text

    section_patterns = {
        "experience": r"(?i)(professional\s+)?experience|work\s+history|employment",
        "education": r"(?i)education|academic\s+background",
        "skills": r"(?i)(technical\s+)?skills|technologies|competencies",
        "projects": r"(?i)projects?|portfolio",
        "certifications": r"(?i)certifications?|licenses",
        "summary": r"(?i)summary|profile|objective",
    }

    section_matches = []
    for section_name, pattern in section_patterns.items():
        for match in re.finditer(pattern, text):
            section_matches.append((match.start(), section_name, match.group(0)))

    section_matches.sort(key=lambda x: x[0])

    for i, (start_pos, section_name, header) in enumerate(section_matches):
        end_pos = section_matches[i + 1][0] if i + 1 < len(section_matches) else len(text)
        section_text = text[start_pos:end_pos].strip()
        result.sections[section_name] = section_text


def _extract_experiences(result: DOCXParseResult) -> None:
    """Extract work experience entries."""
    experience_text = result.sections.get("experience", "")
    if not experience_text:
        return

    lines = experience_text.split("\n")
    current_exp: Optional[Dict[str, str]] = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Date pattern
        date_pattern = r"\d{1,2}/\d{4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}"
        has_date = bool(re.search(date_pattern, line, re.IGNORECASE))

        if has_date and "|" not in line:
            if current_exp:
                result.experiences.append(current_exp)

            current_exp = {
                "raw_header": line,
                "bullets": [],
            }

        elif current_exp and (line.startswith("•") or line.startswith("-") or line.startswith("*")):
            bullet = line.lstrip("•-* ").strip()
            if bullet:
                current_exp["bullets"].append(bullet)

    if current_exp:
        result.experiences.append(current_exp)


def _extract_education(result: DOCXParseResult) -> None:
    """Extract education entries."""
    education_text = result.sections.get("education", "")
    if not education_text:
        return

    lines = [line.strip() for line in education_text.split("\n") if line.strip()]
    current_edu: Optional[Dict[str, str]] = None

    for line in lines:
        degree_keywords = ["bachelor", "master", "phd", "associate", "b.s.", "m.s.", "b.a.", "m.a."]
        has_degree = any(keyword in line.lower() for keyword in degree_keywords)

        university_keywords = ["university", "college", "institute", "school"]
        has_university = any(keyword in line.lower() for keyword in university_keywords)

        if has_degree or has_university:
            if current_edu:
                result.education.append(current_edu)

            current_edu = {
                "raw_line": line,
                "details": [],
            }

        elif current_edu and (line.startswith("•") or line.startswith("-") or "GPA" in line):
            detail = line.lstrip("•-* ").strip()
            if detail:
                current_edu["details"].append(detail)

    if current_edu:
        result.education.append(current_edu)


def _extract_skills(result: DOCXParseResult) -> None:
    """Extract skills section."""
    skills_text = result.sections.get("skills", "")
    if not skills_text:
        return

    lines = skills_text.split("\n")[1:]  # Skip header

    for line in lines:
        line = line.strip()
        if not line or line.lower().startswith("skills"):
            continue

        if "|" in line:
            skills = [s.strip() for s in line.split("|")]
        elif "," in line:
            skills = [s.strip() for s in line.split(",")]
        elif "•" in line or "–" in line:
            skills = [s.strip() for s in re.split(r"[•–]", line)]
        else:
            skills = [line]

        result.skills.extend([s for s in skills if s and not s.endswith(":")])


def _calculate_confidence(result: DOCXParseResult) -> float:
    """Calculate confidence score."""
    score = 0.0

    if result.contact_info.get("name"):
        score += 10
    if result.contact_info.get("email"):
        score += 10
    if result.contact_info.get("phone"):
        score += 5
    if result.contact_info.get("linkedin"):
        score += 5

    if result.experiences:
        score += min(40, len(result.experiences) * 10)

    if result.education:
        score += min(20, len(result.education) * 10)

    if result.skills:
        score += 10

    return min(100.0, score) / 100.0
