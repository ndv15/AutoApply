"""
ATS-optimized resume generator with consistent formatting.
Single-column layout with clean section breaks and proper spacing.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import json


def setup_styles(doc):
    """Configure document-wide styles for consistency."""
    try:
        # Name style (use unique name to avoid conflicts)
        if 'ResumeName' not in doc.styles:
            name_style = doc.styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.name = 'Calibri'
            name_style.font.size = Pt(18)
            name_style.font.bold = True
            name_style.paragraph_format.space_after = Pt(6)

        # Contact style (use unique name to avoid conflicts)
        if 'ResumeContact' not in doc.styles:
            contact_style = doc.styles.add_style('ResumeContact', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.name = 'Calibri'
            contact_style.font.size = Pt(11)
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section heading style
        if 'SectionHeading' not in doc.styles:
            heading_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

        # Body text style (use unique name to avoid conflicts)
        if 'ResumeBodyText' not in doc.styles:
            body_style = doc.styles.add_style('ResumeBodyText', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Calibri'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Bullet style (use unique name to avoid conflicts)
        if 'ResumeBullet' not in doc.styles:
            bullet_style = doc.styles.add_style('ResumeBullet', WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = 'Calibri'
            bullet_style.font.size = Pt(11)
            bullet_style.paragraph_format.left_indent = Inches(0.25)
            bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
            bullet_style.paragraph_format.space_after = Pt(3)
    except:
        # If styles already exist, just continue
        pass

    return doc


def format_role_header(title, company, dates, location, awards=None):
    """Format role header with consistent layout."""
    parts = [f"{title} | {company}"]
    if dates:
        parts.append(dates)
    if location:
        parts.append(location)
    header = " | ".join(parts)
    if awards:
        header = f"{header}\n🏆 {awards}"
    return header


def render_resume(path: Path, identity, summary, soq, experience, skills):
    """
    Generate ATS-optimized resume matching Nate Velasco's template format.

    Args:
        path: Output file path
        identity: Dict with name, contact info
        summary: Executive summary text
        soq: List of qualification statements
        experience: Dict of role_id -> bullets (only APPROVED used)
        skills: List of approved skills
    """
    doc = Document()
    doc = setup_styles(doc)

    # Set margins - tighter for professional look
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header - Name centered
    doc.add_paragraph(identity['name'], style='ResumeName').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact info - centered on two lines
    location = identity.get('location', 'Austin, TX 78725')
    phone = identity.get('phone', '')
    email = identity.get('email', '')
    contact_line = f"{location} | {phone} | {email}"
    doc.add_paragraph(contact_line, style='ResumeContact')

    # LinkedIn URL centered
    linkedin = identity.get('linkedin', '').strip()
    if linkedin:
        doc.add_paragraph(linkedin, style='ResumeContact')

    # Executive Summary
    doc.add_paragraph('Executive Summary', style='SectionHeading')
    doc.add_paragraph(summary, style='ResumeBodyText')

    # Accomplishments section (awards with bullet points)
    awards_path = Path("profile/awards.json")
    if awards_path.exists():
        awards = json.loads(awards_path.read_text())
        if awards:
            doc.add_paragraph('Accomplishments', style='SectionHeading')
            for award in awards:
                # Format: "Award Title (Company) - Year (Level)"
                text = f"➤ {award['title']}"
                details = []
                if award.get('company'):
                    details.append(award['company'])
                if award.get('year'):
                    details.append(str(award['year']))
                if award.get('level'):
                    details.append(award['level'])
                if details:
                    text += f" ({') - ('.join(details)})"
                doc.add_paragraph(text, style='ResumeBullet')

    # Professional Experience
    doc.add_paragraph('Professional Experience', style='SectionHeading')

    # Load full history for dates/locations
    history_path = Path("profile/employment_history.json")
    history = json.loads(history_path.read_text())
    history_map = {h['role_id']: h for h in history}

    # Process experience in order from employment history
    for role in history:
        role_id = role['role_id']
        if role_id in experience and experience[role_id]:
            # Company name header
            company_name = role['company']
            doc.add_paragraph(company_name, style='ResumeBodyText').runs[0].bold = True

            # Title line: "Title" + dates + location
            title = role['title']
            dates = f"{role['start']} – {role.get('end', 'Present')}"
            location = role.get('location', 'Austin, TX (Remote)')

            title_para = doc.add_paragraph(style='ResumeBodyText')
            title_para.add_run(title).bold = True
            title_para.add_run(f" | {dates} | {location}")

            # Awards line if present
            if role.get('awards'):
                awards_text = " | ".join(role['awards'])
                award_para = doc.add_paragraph(style='ResumeBodyText')
                award_para.add_run(f"🏆 {awards_text}").italic = True

            # Bullets - only APPROVED
            for b in experience[role_id]:
                if b.get('status') == 'APPROVED':
                    doc.add_paragraph(f"• {b['text']}", style='ResumeBullet')

            # Add spacing after role
            doc.add_paragraph()

    # Education
    education_path = Path("profile/education.json")
    if education_path.exists():
        education = json.loads(education_path.read_text())
        if education:
            doc.add_paragraph('Education', style='SectionHeading')
            for edu in education:
                degree_text = f"{edu['degree']}"
                if edu.get('major'):
                    degree_text += f" in {edu['major']}"
                if edu.get('concentration'):
                    degree_text += f" with a {edu['concentration']}"
                doc.add_paragraph(degree_text, style='ResumeBodyText')

                # School, location, grad year
                school_parts = []
                if edu.get('school'):
                    school_parts.append(edu['school'])
                if edu.get('location'):
                    school_parts.append(edu['location'])
                if edu.get('grad'):
                    school_parts.append(f"Class of {edu['grad']}")
                if school_parts:
                    doc.add_paragraph(" | ".join(school_parts), style='ResumeBodyText')

                # Highlights
                if edu.get('highlights'):
                    for highlight in edu['highlights']:
                        doc.add_paragraph(f"○ {highlight}", style='ResumeBullet')

    # Summary of Qualifications (from soq parameter)
    if soq:
        doc.add_paragraph('Summary of Qualifications', style='SectionHeading')
        for qual in soq:
            doc.add_paragraph(f"○ {qual}", style='ResumeBullet')

    # Skills - STRICT 3-4 line pipe format using skills_engine formatter
    if skills:
        doc.add_paragraph('Skills', style='SectionHeading')

        # Import formatter
        from ..writing.skills_engine import format_skills_for_resume

        # Get formatted lines (max 4 categories)
        skill_lines = format_skills_for_resume(skills, max_categories=4)

        # Render each line with bold category
        for line in skill_lines:
            # Parse "Category: item | item"
            if ':' in line:
                category, items = line.split(':', 1)
                para = doc.add_paragraph(style='ResumeBodyText')
                para.add_run(f"{category.strip()}: ").bold = True
                para.add_run(items.strip())
            else:
                doc.add_paragraph(line, style='ResumeBodyText')

    doc.save(path)
    return path


def render_resume_preview(
    identity: dict,
    exec_summary: str,
    soq: list,
    approved_bullets: dict,
    skills: list
) -> str:
    """
    Generate HTML preview of resume for live preview panel.

    Args:
        identity: Dict with name, email, phone, location
        exec_summary: Executive summary text
        soq: Summary of qualifications (list of strings)
        approved_bullets: Dict of role_id -> list of bullet dicts
        skills: List of approved skill strings

    Returns:
        HTML string with inline styles matching review_sequential.html
    """
    # Load employment history for role metadata
    history_path = Path("profile/employment_history.json")
    if history_path.exists():
        history = json.loads(history_path.read_text())
        history_map = {h['role_id']: h for h in history}
    else:
        history_map = {}

    html_parts = []

    # Header
    html_parts.append(f'<div class="resume-name">{identity.get("name", "")}</div>')

    # Contact info
    contact_parts = []
    if identity.get("email"):
        contact_parts.append(identity["email"])
    if identity.get("phone"):
        contact_parts.append(identity["phone"])
    if identity.get("location"):
        contact_parts.append(identity["location"])
    html_parts.append(f'<div class="resume-contact">{" | ".join(contact_parts)}</div>')

    # Executive Summary
    html_parts.append('<div class="resume-section">')
    html_parts.append('  <div class="resume-section-title">Executive Summary</div>')
    html_parts.append(f'  <p style="line-height:1.4">{exec_summary}</p>')
    html_parts.append('</div>')

    # Professional Experience
    html_parts.append('<div class="resume-section">')
    html_parts.append('  <div class="resume-section-title">Professional Experience</div>')

    for role_id, bullets in approved_bullets.items():
        if not bullets:
            continue

        role = history_map.get(role_id, {})
        html_parts.append('  <div style="margin-bottom:20px">')

        # Role header
        title = role.get("title", "Unknown Role")
        company = role.get("company", "Unknown Company")
        dates = f"{role.get('start', '')} – {role.get('end', 'Present')}"
        html_parts.append(f'    <div class="resume-role-header">{title} | {company}</div>')
        html_parts.append(f'    <div style="font-size:10pt;color:#666;margin-bottom:8px">{dates}</div>')

        # Bullets container
        html_parts.append(f'    <div id="bullets-{role_id}">')
        for bullet in bullets:
            bullet_text = bullet.get("text", "") if isinstance(bullet, dict) else str(bullet)
            html_parts.append(f'      <div class="resume-bullet">{bullet_text}</div>')
        html_parts.append('    </div>')
        html_parts.append('  </div>')

    html_parts.append('</div>')

    # Skills section
    if skills:
        html_parts.append('<div class="resume-section">')
        html_parts.append('  <div class="resume-section-title">Skills</div>')
        skills_text = " | ".join(skills[:12])  # Max 12 for preview
        html_parts.append(f'  <p style="line-height:1.6">{skills_text}</p>')
        html_parts.append('</div>')

    return '\n'.join(html_parts)
