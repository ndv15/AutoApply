"""
Cover letter generator with company/role customization.
Three-paragraph format: hook → alignment + wins → close.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import json
import datetime


def setup_styles(doc):
    """Configure document styles for consistent formatting."""
    # Check if styles already exist before adding
    try:
        # Header style
        if 'CoverHeader' not in doc.styles:
            header_style = doc.styles.add_style('CoverHeader', WD_STYLE_TYPE.PARAGRAPH)
            header_style.font.name = 'Calibri'
            header_style.font.size = Pt(11)

        # Body style
        if 'CoverBody' not in doc.styles:
            body_style = doc.styles.add_style('CoverBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Calibri'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(12)

        # Signature style
        if 'CoverSignature' not in doc.styles:
            signature_style = doc.styles.add_style('CoverSignature', WD_STYLE_TYPE.PARAGRAPH)
            signature_style.font.name = 'Calibri'
            signature_style.font.size = Pt(11)
            signature_style.font.bold = True
    except:
        # If styles already exist, just continue
        pass

    return doc


def format_date():
    """Format today's date for the letter header."""
    return datetime.datetime.now().strftime("%B %d, %Y")


def extract_top_wins(experience):
    """Extract top quantified wins from approved experience bullets."""
    wins = []
    for role_bullets in experience.values():
        for bullet in role_bullets:
            if bullet.get('status') == 'APPROVED' and any(c.isdigit() for c in bullet['text']):
                wins.append(bullet['text'])
    return sorted(wins, key=lambda x: sum(c.isdigit() for c in x), reverse=True)[:3]


def render_cover(path: Path, company: str, role: str, summary: str, jd: str = "", experience: dict = None):
    """
    Generate customized cover letter with company/role focus.
    
    Args:
        path: Output file path
        company: Company name
        role: Job title
        summary: Executive summary (used for alignment paragraph)
        jd: Job description text (optional, used for customization)
        experience: Dict of role_id -> bullets (for win examples)
    """
    doc = Document()
    doc = setup_styles(doc)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Load identity info
    identity = json.loads(Path("profile/locked_identity.json").read_text())
    
    # Header block
    doc.add_paragraph(format_date(), style='CoverHeader')
    doc.add_paragraph("", style='CoverHeader')  # Spacing
    doc.add_paragraph(f"Hiring Team\n{company}", style='CoverHeader')
    doc.add_paragraph("", style='CoverHeader')  # Spacing

    # Salutation
    doc.add_paragraph(f"Dear Hiring Team,", style='CoverBody')

    # Opening/Hook paragraph
    hook = (
        f"I am writing to express my strong interest in the {role} position at {company}. "
        "With over 7 years of progressive experience in enterprise software sales and a "
        "consistent track record of exceeding quotas and building strategic customer relationships, "
        "I am confident in my ability to make an immediate impact on your team."
    )
    doc.add_paragraph(hook, style='CoverBody')

    # Core alignment + wins
    alignment = summary.split('.')[0] + '.'  # Use first sentence of summary

    # Extract wins if experience provided
    wins_para = ""
    if experience:
        top_wins = extract_top_wins(experience)[:2]  # Get top 2 wins
        if top_wins:
            wins_text = " Furthermore, I have " + " and ".join(
                w.lower().replace('•', '').strip() for w in top_wins
            )
            wins_para = f"\n\n{wins_text}."

    doc.add_paragraph(f"{alignment}{wins_para}", style='CoverBody')

    # Closing paragraph
    close = (
        f"I am excited about the opportunity to contribute to {company}'s continued success "
        "and would welcome the chance to discuss how my background aligns with your needs. "
        "Thank you for considering my application."
    )
    doc.add_paragraph(close, style='CoverBody')

    # Signature
    doc.add_paragraph("Sincerely,", style='CoverBody')
    doc.add_paragraph(identity['name'], style='CoverSignature')

    # Contact block
    contact = (
        f"{identity['email']}\n"
        f"{identity['phone']}\n"
        f"{identity.get('linkedin', '')}"
    )
    doc.add_paragraph(contact, style='CoverHeader')
    
    doc.save(path)
    return path
